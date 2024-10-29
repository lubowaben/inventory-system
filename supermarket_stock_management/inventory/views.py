from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import logout, login as auth_login, authenticate
from django.contrib import messages
from django.db import IntegrityError
from django.utils import timezone
from django.contrib.auth.models import User
from django.http import HttpResponseForbidden,HttpResponse
from django.contrib.auth.forms import PasswordChangeForm
from django.contrib.auth import update_session_auth_hash
from .models import Product, Supplier, StockTransaction, Profile,Cashier,Notification
from .forms import ProductForm, SupplierForm, UserRegisterForm, ProfileUpdateForm, UserUpdateForm,NotificationForm,CashierForm,StockTransactionForm
import pandas as pd
from django.db.models import Sum




# Helper function to check if the user is a manager
def manager_required(view_func):
    def wrapper_func(request, *args, **kwargs):
        if request.user.profile.role != 'manager':
            return HttpResponseForbidden()
        return view_func(request, *args, **kwargs)
    return wrapper_func
def is_manager(user):
    return user.profile.role == 'manager'

# Dashboard view
@login_required
def dashboard(request):
    profile = Profile.objects.get(user=request.user)
    
    if profile.role == 'manager':
        return render(request, 'inventory/dashboard.html')
    
    elif profile.role == 'cashier':
        # Fetch products to display in the cashier dashboard
        products = Product.objects.all()
        context = {
            'products': products
        }
        return render(request, 'inventory/cashier_dashboard.html', context)
    
    else:
        return redirect('login')  # Or any other page if the role is not defined


# Product views
@login_required
def product_list(request):
    products = Product.objects.all()
    return render(request, 'inventory/product_list.html', {'products': products})

@login_required
def add_product(request):
    if request.method == 'POST':
        form = ProductForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Product added successfully!')
            return redirect('product-list')
    else:
        form = ProductForm()
    return render(request, 'inventory/add_product.html', {'form': form})

@login_required
def edit_product(request, pk):
    product = get_object_or_404(Product, pk=pk)
    if request.method == 'POST':
        form = ProductForm(request.POST, instance=product)
        if form.is_valid():
            form.save()
            messages.success(request, 'Product updated successfully!')
            return redirect('product-list')
    else:
        form = ProductForm(instance=product)
    return render(request, 'inventory/edit_product.html', {'form': form})

@login_required
def delete_product(request, pk):
    product = get_object_or_404(Product, pk=pk)
    if request.method == 'POST':
        product.delete()
        messages.success(request, 'Product deleted successfully!')
        return redirect('product-list')
    return render(request, 'inventory/delete_product.html', {'product': product})

# Supplier views
@manager_required
@login_required
def supplier_list(request):
    suppliers = Supplier.objects.all()
    return render(request, 'inventory/supplier_list.html', {'suppliers': suppliers})

@login_required
@manager_required
def add_supplier(request):
    if request.method == 'POST':
        form = SupplierForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Supplier added successfully!')
            return redirect('supplier-list')
    else:
        form = SupplierForm()
    return render(request, 'inventory/add_supplier.html', {'form': form})

@login_required
@manager_required
def edit_supplier(request, pk):
    supplier = get_object_or_404(Supplier, pk=pk)
    if request.method == 'POST':
        form = SupplierForm(request.POST, instance=supplier)
        if form.is_valid():
            form.save()
            messages.success(request, 'Supplier updated successfully!')
            return redirect('supplier-list')
    else:
        form = SupplierForm(instance=supplier)
    return render(request, 'inventory/edit_supplier.html', {'form': form})

@login_required
@manager_required
def delete_supplier(request, pk):
    supplier = get_object_or_404(Supplier, pk=pk)
    if request.method == 'POST':
        supplier.delete()
        messages.success(request, 'Supplier deleted successfully!')
        return redirect('supplier-list')
    return render(request, 'inventory/delete_supplier.html', {'supplier': supplier})

# Transaction views
@login_required
@manager_required
def transaction_list(request):
    transactions = StockTransaction.objects.all()
    return render(request, 'inventory/transaction_list.html', {'transactions': transactions})

# Authentication views
def register(request):
    if request.method == 'POST':
        form = UserRegisterForm(request.POST)
        if form.is_valid():
            user = form.save()
            username = form.cleaned_data.get('username')
            messages.success(request, f'Account created for {username}!')
            auth_login(request, user)
            return redirect('login')
    else:
        form = UserRegisterForm()
    return render(request, 'inventory/register.html', {'form': form})

def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            auth_login(request, user)
            return redirect('dashboard')  # Redirect to the dashboard after login
        else:
            return render(request, 'inventory/login.html', {'error': 'Invalid credentials'})
    else:
        return render(request, 'inventory/login.html')

@login_required
def logout_view(request):
    logout(request)
    return redirect('login')

@login_required
def profile(request):
    if request.method == 'POST':
        u_form = UserUpdateForm(request.POST, instance=request.user)
        p_form = ProfileUpdateForm(request.POST, request.FILES, instance=request.user.profile)
        if u_form.is_valid() and p_form.is_valid():
            u_form.save()
            p_form.save()
            messages.success(request, 'Your account has been updated!')
            return redirect('profile')
    else:
        u_form = UserUpdateForm(instance=request.user)
        p_form = ProfileUpdateForm(instance=request.user.profile)
    return render(request, 'inventory/profile.html', {'u_form': u_form, 'p_form': p_form})

# Password change view
@login_required
@manager_required
def change_password(request):
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user)
            messages.success(request, 'Your password was successfully updated!')
            return redirect('profile')
        else:
            messages.error(request, 'Please correct the error below.')
    else:
        form = PasswordChangeForm(request.user)
    return render(request, 'inventory/change_password.html', {'form': form})

# Manager views for managing cashiers
@login_required
@manager_required
def manage_cashiers(request):
    cashiers = Profile.objects.filter(role='cashier')
    return render(request, 'inventory/manage_cashiers.html', {'cashiers': cashiers})

@login_required
@manager_required


def add_cashier(request):
    if request.method == 'POST':
        form = CashierForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data.get('email')
            if Cashier.objects.filter(email=email).exists():
                messages.error(request, 'A cashier with this email already exists.')
            else:
                try:
                    form.save()
                    messages.success(request, 'Cashier added successfully.')
                    return redirect('cashiers_list')  # Redirect to a cashier list page or appropriate view
                except IntegrityError:
                    messages.error(request, 'There was an error saving the cashier. Please try again.')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = CashierForm()
    
    return render(request, 'inventory/add_cashier.html', {'form': form})



@login_required
@manager_required
def edit_cashier(request, pk):
    user = get_object_or_404(User, pk=pk)
    
    if request.method == 'POST':
        form = UserUpdateForm(request.POST, instance=user)
        if form.is_valid():
            try:
                form.save()
                messages.success(request, 'Cashier updated successfully!')
                return redirect('manage-cashiers')
            except KeyError as e:
                messages.error(request, f'An error occurred: {e}')
                return render(request, 'inventory/edit_cashier.html', {'form': form, 'cashier': user})
    else:
        form = UserUpdateForm(instance=user)
    
    return render(request, 'inventory/edit_cashier.html', {'form': form, 'cashier': user})
@login_required
@manager_required
def delete_cashier(request, pk):
    user = get_object_or_404(User, pk=pk)
    if request.method == 'POST':
        user.delete()
        messages.success(request, 'Cashier deleted successfully!')
        return redirect('manage-cashiers')
    return render(request, 'inventory/delete_cashier.html', {'cashier': user})
@login_required
def create_profile(request):
    # Ensure that a profile is created only if it doesn't already exist
    if not hasattr(request.user, 'profile'):
        Profile.objects.create(user=request.user)
        messages.success(request, 'Profile created successfully!')
    else:
        messages.info(request, 'Profile already exists.')
    return redirect('profile')  # Redirect to the profile page or wherever appropriate
def add_notification(request):
    if request.method == 'POST':
        form = NotificationForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('dashboard')  # Redirect to the dashboard or another appropriate page
    else:
        form = NotificationForm()
    
    return render(request, 'inventory/add_notification.html', {'form': form})
@login_required
@manager_required
def cashier_list(request):
    cashiers = Cashier.objects.all()
    return render(request, 'inventory/cashier_list.html', {'cashiers': cashiers})




def add_transaction(request):
    if request.method == 'POST':
        form = StockTransactionForm(request.POST)
        if form.is_valid():
            transaction = form.save(commit=False)
            transaction.date = timezone.now()  # Set the date here
            transaction.save()
            return redirect('dashboard')  # Redirect to the dashboard or another appropriate page
    else:
        form = StockTransactionForm()
    
    
    return render(request, 'inventory/add_transaction.html', {'form': form})

def transaction_list(request):
    transactions = StockTransaction.objects.all()
    return render(request, 'inventory/transaction_list.html', {'transactions': transactions})
def edit_transaction(request, transaction_id):
    transaction = get_object_or_404(StockTransaction, id=transaction_id)

    if request.method == 'POST':
        form = StockTransactionForm(request.POST, instance=transaction)
        if form.is_valid():
            form.save()
            return redirect('dashboard')  # Redirect to a relevant page after saving
    else:
        form = StockTransactionForm(instance=transaction)

    return render(request, 'inventory/edit_transaction.html', {'form': form, 'transaction': transaction})

@login_required
@manager_required
def delete_transaction(request, transaction_id):
    transaction = get_object_or_404(StockTransaction, id=transaction_id)

    if request.method == 'POST':
        transaction.delete()
        return redirect('dashboard')  # Redirect to a relevant page after deleting

    return render(request, 'inventory/delete_transaction.html', {'transaction': transaction})
@login_required
@manager_required

def add_cashier(request):
    if request.method == 'POST':
        form = CashierForm(request.POST)
        if form.is_valid():
            # Save the form data to the database
            form.save()
            messages.success(request, 'Cashier added successfully.')
            return redirect('cashier_list')  # Redirect to a page listing all cashiers or another page
        else:
            # Handle form errors
            messages.error(request, 'Please correct the errors below.')
    else:
        form = CashierForm()
    
    return render(request, 'inventory/add_cashier.html', {'form': form})

def process_sale(request):
    if request.method == 'POST':
        product_id = request.POST.get('product_id')
        quantity_sold = request.POST.get('quantity')
        
        # Check for missing product_id or quantity
        if not product_id or not quantity_sold:
            messages.error(request, 'Product ID or quantity is missing.')
            return redirect('product-list')

        try:
            # Convert quantity_sold to an integer
            quantity_sold = int(quantity_sold)
        except ValueError:
            messages.error(request, 'Quantity must be a number.')
            return redirect('product-list')

        try:
            # Fetch the product
            product = Product.objects.get(id=product_id)
        except Product.DoesNotExist:
            messages.error(request, 'Product does not exist.')
            return redirect('product-list')

        # Check for sufficient stock
        if product.quantity < quantity_sold:
            messages.error(request, 'Insufficient stock.')
            return redirect('product-list')

        # Process the sale
        product.quantity -= quantity_sold
        product.save()

        # Record the transaction
        StockTransaction.objects.create(
            product=product,
            quantity=quantity_sold,
            transaction_type='sale',
            date=timezone.now()
        )

        messages.success(request, 'Sale processed successfully.')
        return redirect('dashboard')
    
    messages.error(request, 'Invalid request method.')
    return redirect('dashboard')


def edit_notification(request, id):
    notification = get_object_or_404(Notification, id=id)
    # Your logic here
    return render(request, 'inventory/edit_notification.html', {'notification': notification})
@login_required
@manager_required
def delete_notification(request, notification_id):
    notification = get_object_or_404(Notification, id=notification_id)

    if request.method == 'POST':
        notification.delete()
        messages.success(request, 'Notification deleted successfully.')
        return redirect('dashboard')  # Redirect to the dashboard or wherever you need after deletion

    return render(request, 'delete_notification.html', {'notification': notification})
def save(self, *args, **kwargs):
    if self.pk is None:  # Check if the transaction is being created, not updated
        if self.transaction_type == 'sale':
            if self.product.quantity >= self.quantity:
                self.product.quantity -= self.quantity
            else:
                raise ValueError('Not enough stock to complete the sale.')
        elif self.transaction_type == 'restock':
            self.product.quantity += self.quantity

        self.product.save()  # Save the updated product quantity

    super().save(*args, **kwargs)  # Call the original save method
def notification_list(request):
    notifications = Notification.objects.all().order_by('-created_at')
    return render(request, 'inventory/notification_list.html', {'notifications': notifications})
def calculate_totals():
    total_sales = StockTransaction.objects.filter(transaction_type='sale').aggregate(total=Sum('quantity'))['total']
    total_stock = Product.objects.aggregate(total=Sum('quantity'))['total']
    total_transactions = StockTransaction.objects.aggregate(total=Sum('quantity'))['total']
    return total_sales, total_stock, total_transactions


# Generate Excel report
def export_excel(request):
    total_sales, total_stock, total_transactions = calculate_totals()

    # Create a DataFrame
    data = {
        'Total Sales': [total_sales],
        'Total Stock': [total_stock],
        'Total Transactions': [total_transactions],
    }
    df = pd.DataFrame(data)

    # Create a response object with the correct MIME type
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=report.xlsx'

    # Write the DataFrame to an Excel file
    df.to_excel(response, index=False)

    return response

# Generate Word report
def export_word(request):
    from docx import Document

    total_sales, total_stock, total_transactions = calculate_totals()

    # Create a Word document
    doc = Document()
    doc.add_heading('Sales and Stock Report', 0)

    doc.add_paragraph(f'Total Sales: {total_sales}')
    doc.add_paragraph(f'Total Stock: {total_stock}')
    doc.add_paragraph(f'Total Transactions: {total_transactions}')

    # Create a response object with the correct MIME type
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=report.docx'

    # Save the document to the response
    doc.save(response)

    return response
import csv


def stock_report(request):
    products = Product.objects.all()
    
    # Check if the export to CSV option is requested
    if 'export' in request.GET:
        # Create the HttpResponse object with the appropriate CSV header.
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="stock_report.csv"'

        writer = csv.writer(response)
        writer.writerow(['Item Name', 'Quantity Available', 'Low Stock Signal'])

        for product in products:
            low_stock_signal = 'Yes' if product.quantity < 100 else 'No'
            writer.writerow([product.name, product.quantity, low_stock_signal])

        return response

    # Adding a signal flag for items with quantity below 100 for display in template
    for product in products:
        if product.quantity < 100:
            product.low_stock = True
        else:
            product.low_stock = False

    context = {
        'products': products,
    }
    
    return render(request, 'inventory/stock_report.html', context)
